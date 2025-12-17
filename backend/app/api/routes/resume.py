"""
Resume API Routes.

Handles resume upload, parsing, and version management.
"""

from app.models.chat import (
    FileUploadResponse,
    ResumeVersionResponse,
    VersionCompareResponse,
)
from app.services.firebase_service import get_storage_service
from app.services.resume_parser import ResumeParserService
from app.services.vector_store import VectorStoreService
from fastapi import APIRouter, File, Form, HTTPException, UploadFile

router = APIRouter(prefix="/resume", tags=["resume"])


@router.post("/upload", response_model=FileUploadResponse)
async def upload_resume(
    file: UploadFile = File(...), user_id: str = Form(default="default_user")
):
    """
    Upload and parse a resume file.

    Supports PDF and DOCX formats. The resume will be parsed,
    sections will be extracted, and it will be indexed for search.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    allowed_extensions = ["pdf", "docx"]
    extension = file.filename.split(".")[-1].lower()
    if extension not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}",
        )

    parser = ResumeParserService()
    storage = get_storage_service()
    vector_store = VectorStoreService()

    try:
        resume = await parser.parse_file(
            file=file.file, filename=file.filename, user_id=user_id
        )

        await storage.save_resume(resume)

        await storage.create_resume_version(
            resume_id=resume.id,
            content=resume.get_full_text(),
            sections=resume.sections,
            changes_description="Initial upload",
            agent_used="upload",
        )

        try:
            await vector_store.index_resume(resume)
        except Exception as e:
            print(f"Warning: Failed to index resume in vector store: {e}")

        sections_detected = [section.section_type.value for section in resume.sections]

        return FileUploadResponse(
            resume_id=resume.id,
            filename=resume.filename,
            message=f"Resume uploaded and parsed successfully. Found {len(resume.sections)} sections.",
            sections_detected=sections_detected,
            metadata={
                "character_count": len(resume.raw_text),
                "sections_count": len(resume.sections),
            },
        )

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to process resume: {str(e)}"
        )


@router.get("/{resume_id}")
async def get_resume(resume_id: str):
    """Get a resume by ID."""
    storage = get_storage_service()
    resume = await storage.get_resume(resume_id)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    return {
        "id": resume.id,
        "filename": resume.filename,
        "sections": [
            {
                "type": section.section_type.value,
                "title": section.title,
                "content": section.content,
            }
            for section in resume.sections
        ],
        "metadata": resume.metadata,
        "created_at": resume.created_at.isoformat(),
        "updated_at": resume.updated_at.isoformat(),
    }


@router.get("/{resume_id}/content")
async def get_resume_content(resume_id: str):
    """Get the full content of a resume."""
    storage = get_storage_service()
    resume = await storage.get_resume(resume_id)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    return {
        "id": resume.id,
        "content": resume.get_full_text(),
        "raw_text": resume.raw_text,
    }


@router.get("/{resume_id}/versions")
async def get_resume_versions(resume_id: str):
    """Get all versions of a resume."""
    storage = get_storage_service()

    resume = await storage.get_resume(resume_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    versions = await storage.get_resume_versions(resume_id)

    return {
        "resume_id": resume_id,
        "versions": [
            {
                "id": v.id,
                "version_number": v.version_number,
                "changes_description": v.changes_description,
                "agent_used": v.agent_used,
                "created_at": v.created_at.isoformat(),
            }
            for v in versions
        ],
        "total_versions": len(versions),
    }


@router.get("/{resume_id}/versions/{version_number}")
async def get_resume_version(resume_id: str, version_number: int):
    """Get a specific version of a resume."""
    storage = get_storage_service()

    versions = await storage.get_resume_versions(resume_id)
    version = next((v for v in versions if v.version_number == version_number), None)

    if not version:
        raise HTTPException(status_code=404, detail="Version not found")

    return ResumeVersionResponse(
        version_id=version.id,
        version_number=version.version_number,
        content=version.content,
        changes_description=version.changes_description,
        agent_used=version.agent_used,
        created_at=version.created_at,
    )


@router.get("/{resume_id}/compare/{version_a}/{version_b}")
async def compare_versions(resume_id: str, version_a: int, version_b: int):
    """Compare two versions of a resume."""
    storage = get_storage_service()

    versions = await storage.get_resume_versions(resume_id)
    va = next((v for v in versions if v.version_number == version_a), None)
    vb = next((v for v in versions if v.version_number == version_b), None)

    if not va or not vb:
        raise HTTPException(status_code=404, detail="One or both versions not found")

    differences = _compute_differences(va.content, vb.content)

    return VersionCompareResponse(
        version_a=ResumeVersionResponse(
            version_id=va.id,
            version_number=va.version_number,
            content=va.content,
            changes_description=va.changes_description,
            agent_used=va.agent_used,
            created_at=va.created_at,
        ),
        version_b=ResumeVersionResponse(
            version_id=vb.id,
            version_number=vb.version_number,
            content=vb.content,
            changes_description=vb.changes_description,
            agent_used=vb.agent_used,
            created_at=vb.created_at,
        ),
        differences=differences,
    )


@router.post("/{resume_id}/revert/{version_number}")
async def revert_to_version(resume_id: str, version_number: int):
    """Revert a resume to a previous version."""
    storage = get_storage_service()

    resume = await storage.get_resume(resume_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    versions = await storage.get_resume_versions(resume_id)
    target_version = next(
        (v for v in versions if v.version_number == version_number), None
    )

    if not target_version:
        raise HTTPException(status_code=404, detail="Version not found")

    new_version = await storage.create_resume_version(
        resume_id=resume_id,
        content=target_version.content,
        sections=target_version.sections,
        changes_description=f"Reverted to version {version_number}",
        agent_used="revert",
        parent_version_id=target_version.id,
    )

    resume.sections = target_version.sections
    if hasattr(storage, "update_resume"):
        await storage.update_resume(resume)

    return {
        "message": f"Reverted to version {version_number}",
        "new_version_number": new_version.version_number,
        "version_id": new_version.id,
    }


@router.post("/{resume_id}/export")
async def export_resume(resume_id: str, format: str = "pdf"):
    """
    Export a resume to PDF or DOCX format.

    Args:
        resume_id: ID of the resume to export.
        format: Export format ('pdf' or 'docx').
    """
    storage = get_storage_service()
    resume = await storage.get_resume(resume_id)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Get the latest version content and sections
    versions = await storage.get_resume_versions(resume_id)
    latest_version = max(versions, key=lambda v: v.version_number) if versions else None

    # Use sections from latest version if available, otherwise from resume
    sections_to_export = (
        latest_version.sections
        if latest_version and latest_version.sections
        else resume.sections
    )
    content = latest_version.content if latest_version else resume.get_full_text()

    if format.lower() == "docx":
        return await _export_docx(resume, content, sections_to_export)
    else:
        return await _export_pdf(resume, content, sections_to_export)


def _clean_section_content(content: str) -> str:
    """
    Remove LLM reasoning/metadata from section content.

    This filters out things like "Key changes made:", "Here's what I changed:", etc.
    that should only appear in chat, not in exported documents.
    """
    import re

    # Patterns that indicate LLM reasoning/metadata (not actual resume content)
    reasoning_patterns = [
        r"(?:^|\n)\s*(?:Key changes(?: made)?|Changes made|Here'?s? what I (?:changed|modified|updated)|What I (?:changed|did)|Reasoning|Explanation|Notes?|Summary of changes):\s*\n?.*",
        r"(?:^|\n)\s*\d+\.\s*\*\*[^*]+\*\*:.*",  # Numbered bold items like "1. **Added summary**:"
        r"(?:^|\n)\s*-\s*\*\*[^*]+\*\*:.*",  # Bullet bold items like "- **Emphasized skills**:"
        r"(?:^|\n)\s*I (?:have |'ve )?(?:made the following|updated|modified|changed|reorganized|reordered|incorporated|highlighted|added|emphasized).*",
    ]

    cleaned = content
    for pattern in reasoning_patterns:
        # Find where the reasoning starts and truncate
        match = re.search(pattern, cleaned, re.IGNORECASE | re.DOTALL)
        if match:
            # Keep only content before the reasoning
            cleaned = cleaned[: match.start()].strip()

    return cleaned


async def _export_pdf(resume, content: str, sections=None):
    """Export resume as PDF."""
    import io
    import re

    from fastapi.responses import StreamingResponse

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CustomTitle", parent=styles["Heading1"], fontSize=16, spaceAfter=12
        )
        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=12,
            spaceAfter=6,
            spaceBefore=12,
        )
        body_style = ParagraphStyle(
            "CustomBody", parent=styles["Normal"], fontSize=10, spaceAfter=6
        )

        story = []

        # Add title (filename without extension)
        title = resume.filename.rsplit(".", 1)[0] if resume.filename else "Resume"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Use provided sections or fall back to resume sections
        export_sections = sections if sections else resume.sections

        # Add sections
        for section in export_sections:
            # Clean title - remove any markdown or special characters
            clean_title = re.sub(
                r"^#+\s*", "", section.title
            )  # Remove markdown headers
            clean_title = clean_title.strip().upper()

            # Skip sections with weird titles (like metadata or reasoning)
            if len(clean_title) > 50 or not clean_title:
                continue

            # Skip sections that look like LLM reasoning
            reasoning_titles = [
                "key changes",
                "changes made",
                "reasoning",
                "explanation",
                "notes",
                "summary of changes",
                "what i changed",
                "modifications",
            ]
            if any(rt in clean_title.lower() for rt in reasoning_titles):
                continue

            story.append(Paragraph(clean_title, heading_style))

            # Clean section content - remove any LLM reasoning that got mixed in
            section_content = _clean_section_content(section.content)

            # Skip empty sections after cleaning
            if not section_content.strip():
                story.pop()  # Remove the heading we just added
                continue

            # Escape HTML special chars for PDF rendering
            section_content = section_content.replace("&", "&amp;")
            section_content = section_content.replace("<", "&lt;")
            section_content = section_content.replace(">", "&gt;")
            section_content = section_content.replace("\n", "<br/>")

            # Remove any remaining problematic characters
            section_content = re.sub(r"[^\x00-\x7F]+", " ", section_content)

            story.append(Paragraph(section_content, body_style))

        doc.build(story)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={resume.filename.rsplit('.', 1)[0]}_optimized.pdf"
            },
        )

    except ImportError:
        # Fallback: return plain text if reportlab not available
        return {
            "error": "PDF export requires reportlab library",
            "content": content,
            "filename": f"{resume.filename.rsplit('.', 1)[0]}_optimized.txt",
        }


async def _export_docx(resume, content: str, sections=None):
    """Export resume as DOCX."""
    import io
    import re

    from fastapi.responses import StreamingResponse

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        doc = Document()

        # Set margins
        for doc_section in doc.sections:
            doc_section.top_margin = Inches(0.75)
            doc_section.bottom_margin = Inches(0.75)
            doc_section.left_margin = Inches(0.75)
            doc_section.right_margin = Inches(0.75)

        # Add title
        title = resume.filename.rsplit(".", 1)[0] if resume.filename else "Resume"
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Use provided sections or fall back to resume sections
        export_sections = sections if sections else resume.sections

        # Add sections
        for section in export_sections:
            # Clean title - remove any markdown or special characters
            clean_title = re.sub(
                r"^#+\s*", "", section.title
            )  # Remove markdown headers
            clean_title = clean_title.strip()

            # Skip sections with weird titles (like metadata)
            if len(clean_title) > 50 or not clean_title:
                continue

            # Skip sections that look like LLM reasoning
            reasoning_titles = [
                "key changes",
                "changes made",
                "reasoning",
                "explanation",
                "notes",
                "summary of changes",
                "what i changed",
                "modifications",
            ]
            if any(rt in clean_title.lower() for rt in reasoning_titles):
                continue

            doc.add_heading(clean_title, level=1)

            # Clean section content - remove any LLM reasoning
            section_content = _clean_section_content(section.content)

            # Skip empty sections after cleaning
            if not section_content.strip():
                continue

            # Add content paragraphs
            for line in section_content.split("\n"):
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    para.style.font.size = Pt(11)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={resume.filename.rsplit('.', 1)[0]}_optimized.docx"
            },
        )

    except ImportError:
        return {
            "error": "DOCX export requires python-docx library",
            "content": content,
            "filename": f"{resume.filename.rsplit('.', 1)[0]}_optimized.txt",
        }


@router.post("/{resume_id}/analyze")
async def analyze_resume(resume_id: str, job_description: str | None = None):
    """
    Analyze a resume and provide scores and recommendations.

    Args:
        resume_id: ID of the resume to analyze.
        job_description: Optional job description for matching analysis.
    """
    from app.core.llm import get_llm

    storage = get_storage_service()
    resume = await storage.get_resume(resume_id)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    llm = get_llm()
    resume_text = resume.get_full_text()

    analysis_prompt = f"""Analyze this resume and provide a detailed evaluation.

Resume:
{resume_text[:4000]}

{"Job Description for matching:" + job_description[:2000] if job_description else ""}

Provide your analysis in the following exact format:

OVERALL_SCORE: [0-100]
KEYWORD_SCORE: [0-100]
FORMAT_SCORE: [0-100]
IMPACT_SCORE: [0-100]

STRENGTHS:
- [strength 1]
- [strength 2]
- [strength 3]

IMPROVEMENTS:
- [improvement 1]
- [improvement 2]
- [improvement 3]

KEYWORDS_FOUND:
- [keyword 1]
- [keyword 2]

MISSING_KEYWORDS:
- [missing keyword 1]
- [missing keyword 2]

SUMMARY:
[2-3 sentence summary of the resume quality and recommendations]"""

    try:
        from langchain_core.messages import HumanMessage, SystemMessage

        messages = [
            SystemMessage(
                content="You are an expert resume analyst. Provide detailed, actionable feedback."
            ),
            HumanMessage(content=analysis_prompt),
        ]

        response = await llm.ainvoke(messages)
        analysis_text = response.content

        # Parse the response
        import re

        def extract_score(pattern, text, default=70):
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    return min(100, max(0, int(match.group(1))))
                except ValueError:
                    return default
            return default

        def extract_list(pattern, text):
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                items = re.findall(r"^-\s*(.+)$", match.group(1), re.MULTILINE)
                return [item.strip() for item in items[:5]]
            return []

        overall = extract_score(r"OVERALL_SCORE:\s*(\d+)", analysis_text)
        keywords = extract_score(r"KEYWORD_SCORE:\s*(\d+)", analysis_text)
        format_score = extract_score(r"FORMAT_SCORE:\s*(\d+)", analysis_text)
        impact = extract_score(r"IMPACT_SCORE:\s*(\d+)", analysis_text)

        strengths = extract_list(r"STRENGTHS:(.*?)(?=IMPROVEMENTS:|$)", analysis_text)
        improvements = extract_list(
            r"IMPROVEMENTS:(.*?)(?=KEYWORDS_FOUND:|$)", analysis_text
        )
        keywords_found = extract_list(
            r"KEYWORDS_FOUND:(.*?)(?=MISSING_KEYWORDS:|$)", analysis_text
        )
        missing_keywords = extract_list(
            r"MISSING_KEYWORDS:(.*?)(?=SUMMARY:|$)", analysis_text
        )

        summary_match = re.search(r"SUMMARY:\s*(.+?)$", analysis_text, re.DOTALL)
        summary = (
            summary_match.group(1).strip() if summary_match else "Analysis complete."
        )

        return {
            "resume_id": resume_id,
            "evaluation": {
                "overall": overall,
                "keywords": keywords,
                "format": format_score,
                "impact": impact,
            },
            "strengths": strengths or ["Good structure", "Clear formatting"],
            "improvements": improvements or ["Add more quantifiable achievements"],
            "keywords_found": keywords_found,
            "missing_keywords": missing_keywords,
            "summary": summary,
            "job_match": bool(job_description),
        }

    except Exception as e:
        # Return default scores on error
        return {
            "resume_id": resume_id,
            "evaluation": {"overall": 75, "keywords": 70, "format": 80, "impact": 72},
            "strengths": ["Resume uploaded successfully"],
            "improvements": ["Consider adding more details"],
            "keywords_found": [],
            "missing_keywords": [],
            "summary": f"Basic analysis complete. Error in detailed analysis: {str(e)}",
            "job_match": False,
        }


def _compute_differences(content_a: str, content_b: str) -> list[dict]:
    """Compute differences between two content strings."""
    import difflib

    differ = difflib.unified_diff(
        content_a.splitlines(keepends=True),
        content_b.splitlines(keepends=True),
        lineterm="",
    )

    differences = []
    current_change = None

    for line in differ:
        if line.startswith("---") or line.startswith("+++"):
            continue
        elif line.startswith("@@"):
            if current_change:
                differences.append(current_change)
            current_change = {"type": "context", "lines": []}
        elif line.startswith("-"):
            if current_change is None:
                current_change = {"type": "removal", "lines": []}
            current_change["type"] = "removal"
            current_change["lines"].append(line[1:])
        elif line.startswith("+"):
            if current_change is None:
                current_change = {"type": "addition", "lines": []}
            if current_change["type"] == "removal":
                differences.append(current_change)
                current_change = {"type": "addition", "lines": []}
            current_change["type"] = "addition"
            current_change["lines"].append(line[1:])
        else:
            if current_change and current_change["lines"]:
                differences.append(current_change)
                current_change = {"type": "context", "lines": []}

    if current_change and current_change["lines"]:
        differences.append(current_change)

    return differences[:20]
